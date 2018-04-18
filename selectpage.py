#!/usr/bin/env python
# -*- encoding: utf-8 -*-

import pymongo
import uuid
from pymongo import MongoClient
from docx import Document
from docx.shared  import Pt


def get_system_db(host, port, system_name):
    client = MongoClient(host, port)
    system_db = client[system_name]
    return system_db

def get_main_db():
    client = MongoClient('192.168.0.102', 27017)
    dev_db = client['devplatform']
    return dev_db

def list_autodesign_data(db):
    collection = db.autodesign2.find({},{"objs":1,"_id":0})
    for item in collection:
        for key, value in item['objs'].iteritems():
            data = value.get('option', {}).get('data')
            if not isinstance(data, list):
                continue
            yield (item, key, data)

def get_rule_name_by_url(url, db):
    rule_name = []
    for item, key, data in  list_autodesign_data(db):
        for data_item in data:
            if not isinstance(data_item.get('data'), dict):
                continue

            if 'url' not in data_item['data']:
                continue

            if data_item['data']['url'] == url:
                key = key.encode('gbk')
                rule_name.append(item['objs'][key]['alias'])

    return rule_name

def get_ruleid_by_rule_name(name):
    dev_db = get_main_db()
    dev_collection = dev_db.rule.find({"alias":name})
    for dev_item in dev_collection:
        if isinstance(dev_item, dict):
            rule_id = dev_item.get("_id")
            return rule_id
        else:
            print("no rule")

def get_page_equal_ruleid(rule_id, system_db):

    gz_collection = system_db.page_setup.find({},{"relations":1})
    for item in gz_collection:
        if not item.has_key("relations"):
            continue
        if not item["relations"]:
            continue

        for key, value in item["relations"]["MDict"].items():
            if not value or value['type'] != 1:
                continue

            for tag_name, tag_value in value['value']['module']['binding'].items():
                relation_ruleid = tag_value['datasource']['id']
                if relation_ruleid.encode('utf-8') == str(rule_id):
                    dev_db = get_main_db()
                    page_collection = dev_db.page.find({"_id":item['_id']})
                    for page_item in page_collection:
                        page =[page_value for page_key,page_value in page_item.items() if page_key == 'alias']

                    return page


class gen_doc():
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.font.size = Pt(24)
    run.font.name = u'宋体'

    def __init__(self, **kwargs):
        self.name = kwargs.pop('name', None)
        self.url = kwargs.pop('url', None)
        self.system = kwargs.pop('system', None)

    def run(self):
        if self.url and self.system:
            if isinstance(self.name, list):
                self.document.add_heading(text=u'与' + self.url + u'相关的规则与页面', level=1)
                for name_item in self.name:
                    rule_id = get_ruleid_by_rule_name(name_item)
                    page = get_page_equal_ruleid(rule_id, self.system)
                    if isinstance(page,list):
                        self.document.add_heading(text=u'用了' + name_item + u'规则的页面', level=2)
                        for item in page:
                            self.document.add_paragraph(text=item, style='ListBullet')
            self.document.save(u'使用了url' + '.docx')

        else:
            if isinstance(self.name, str):
                self.document.add_heading(text=u'用了'+ self.name.decode('utf-8')+u'规则的页面',level=1)
                rule_id = get_ruleid_by_rule_name(self.name)
                page = get_page_equal_ruleid(rule_id, self.system)
                if page and isinstance(page, list):
                    for item in page:
                        self.document.add_paragraph(text=item, style='ListBullet')
            self.document.save(self.name.decode('utf-8') + '.docx')
        #document.add_page_break()


def main():
    url ="http://58.56.160.41:38015/web/search/common/articlewy/select"
    name =""
    child_system_name= "gz_opinion"
    host='192.168.0.102'
    port = 27017
    db = None
    if host and port:
        db = get_system_db(host,port,child_system_name)
    if url and db:
        rule_name = get_rule_name_by_url(url, db)
        get_doc = gen_doc(url=url,name=rule_name,system=db)
        get_doc.run()
    if name and child_system_name:
        get_doc = gen_doc(name=name, system=db)
        get_doc.run()


if __name__ == "__main__":
    main()
